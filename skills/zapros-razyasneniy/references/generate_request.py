"""
Генерация Word-файла запроса на разъяснение положений извещения по 44-ФЗ.

Все вопросы по закупке упаковываются в ОДИН файл — это списывается
с лимита запросов (для электронного аукциона — 3 запроса от одного лица)
один раз.

Реквизиты заказчика и заявителя в файл НЕ включаются — площадка ЕИС
прикрепляет их автоматически при подаче запроса.

Использование:
    python generate_request.py --input data.json --output "Запрос разъяснений по закупке № X.docx"

Где data.json — JSON-файл со структурой:
{
    "zakupka": {
        "reestr_number": "0121600013926000030",
        "request_number": null    // null или 2, 3 если это не первый запрос
    },
    "questions": [
        {
            "topic": "Эквивалент по позиции 4 ЛСР",
            "paragraphs": [
                "Первый абзац вопроса с цитатой...",
                "Второй абзац с конкретным запросом и вариантами ответа..."
            ]
        },
        {
            "topic": "Коллизия по объёму архива",
            "paragraphs": [
                "Текст второго вопроса...",
                "Прошу уточнить..."
            ]
        }
    ],
    "podpisant": {
        "name_short": "И.И. Иванов",
        "date": "02.05.2026"        // null = текущая дата
    }
}

Зависимости: python-docx (pip install python-docx)
"""

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx не установлен. Запустите: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def set_default_font(document: Document) -> None:
    """Установить шрифт Times New Roman 12pt для всего документа."""
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), 'Times New Roman')
    rfonts.set(qn('w:hAnsi'), 'Times New Roman')
    rfonts.set(qn('w:cs'), 'Times New Roman')


def set_page_margins(document: Document) -> None:
    """А4 + поля 2/2/3/1.5 см."""
    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.page_width = Mm(210)
        section.page_height = Mm(297)


def add_paragraph(
    document: Document,
    text: str,
    *,
    alignment: int = WD_ALIGN_PARAGRAPH.JUSTIFY,
    bold: bool = False,
    italic: bool = False,
    size_pt: int = 12,
    line_spacing: float = 1.15,
):
    """Добавить параграф с настройками шрифта и выравнивания."""
    p = document.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), 'Times New Roman')
    rfonts.set(qn('w:hAnsi'), 'Times New Roman')
    rfonts.set(qn('w:cs'), 'Times New Roman')
    return p


def add_blank_line(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)


def render_zagolovok(document: Document, data: dict) -> None:
    """Заголовок документа."""
    add_paragraph(
        document,
        "ЗАПРОС НА РАЗЪЯСНЕНИЕ",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size_pt=14,
    )
    add_paragraph(
        document,
        "положений извещения о проведении закупки",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size_pt=12,
    )

    z = data['zakupka']
    title = f"№ {z['reestr_number']}"
    if z.get('request_number'):
        title = f"Запрос № {z['request_number']} к закупке № {z['reestr_number']}"
    add_paragraph(
        document,
        title,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size_pt=14,
    )


def render_vstuplenie(document: Document) -> None:
    """Обращение + вступительная фраза."""
    add_blank_line(document)
    add_blank_line(document)
    add_paragraph(document, "Уважаемые специалисты!", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_blank_line(document)
    add_paragraph(
        document,
        "Прошу разъяснить положения извещения о проведении закупки "
        "по следующим вопросам.",
    )


def render_voprosy(document: Document, data: dict) -> None:
    """Пронумерованные блоки вопросов."""
    questions = data.get('questions') or []
    if not questions:
        raise ValueError("Список вопросов пуст. Добавьте хотя бы один вопрос в data.questions.")

    for index, question in enumerate(questions, start=1):
        add_blank_line(document)
        add_blank_line(document)

        # Подзаголовок блока
        topic = question.get('topic', '').strip()
        if topic:
            add_paragraph(
                document,
                f"{index}. {topic}",
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                bold=True,
            )
        else:
            add_paragraph(
                document,
                f"{index}.",
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                bold=True,
            )

        add_blank_line(document)

        for paragraph_text in question.get('paragraphs', []):
            add_paragraph(document, paragraph_text)
            add_blank_line(document)


def render_podpis(document: Document, data: dict) -> None:
    """Подпись и дата."""
    add_blank_line(document)
    add_paragraph(document, "С уважением,", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_blank_line(document)

    p = data.get('podpisant') or {}
    name = p.get('name_short', '')
    if name:
        add_paragraph(document, name, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    add_blank_line(document)
    date_str = p.get('date') or datetime.now().strftime('%d.%m.%Y')
    add_paragraph(document, date_str, alignment=WD_ALIGN_PARAGRAPH.LEFT)


def generate_request(data: dict, output_path: Path) -> None:
    """Сгенерировать Word-файл запроса."""
    document = Document()
    set_default_font(document)
    set_page_margins(document)

    render_zagolovok(document, data)
    render_vstuplenie(document)
    render_voprosy(document, data)
    render_podpis(document, data)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def build_default_filename(data: dict) -> str:
    """Шаблон имени файла: 'Запрос разъяснений по закупке № N.docx'"""
    z = data['zakupka']
    if z.get('request_number'):
        return f"Запрос разъяснений № {z['request_number']} по закупке № {z['reestr_number']}.docx"
    return f"Запрос разъяснений по закупке № {z['reestr_number']}.docx"


def main():
    parser = argparse.ArgumentParser(
        description="Генерация Word-файла запроса на разъяснение положений извещения по 44-ФЗ"
    )
    parser.add_argument(
        '--input', '-i',
        type=Path,
        required=True,
        help='Путь к JSON-файлу с входными данными',
    )
    parser.add_argument(
        '--output', '-o',
        type=Path,
        default=None,
        help='Путь к выходному .docx файлу (по умолчанию — авто из реестрового номера)',
    )
    args = parser.parse_args()

    if not args.input.exists():
        print(f"ERROR: входной файл не найден: {args.input}", file=sys.stderr)
        sys.exit(1)

    with args.input.open('r', encoding='utf-8') as f:
        data = json.load(f)

    output_path = args.output
    if output_path is None:
        output_path = Path(build_default_filename(data))

    generate_request(data, output_path)
    print(f"OK: {output_path}")


if __name__ == '__main__':
    main()
