"""Extract text from all documents in a tender folder into a flat text dump."""
import sys
import os
from pathlib import Path

def extract_docx(path):
    from docx import Document
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)

def extract_doc(path):
    """Convert .doc to .docx via Word, then extract."""
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        abs_path = str(Path(path).resolve())
        tmp_docx = abs_path + ".tmp.docx"
        doc = word.Documents.Open(abs_path, ReadOnly=True)
        doc.SaveAs2(tmp_docx, FileFormat=16)  # 16 = wdFormatDocumentDefault (.docx)
        doc.Close(False)
        text = extract_docx(tmp_docx)
        os.remove(tmp_docx)
        return text
    finally:
        word.Quit()

def extract_xlsx(path):
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True)
    parts = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        parts.append(f"\n=== Sheet: {sheet} ===")
        for row in ws.iter_rows(values_only=True):
            row_str = " | ".join(str(c) if c is not None else "" for c in row)
            if row_str.strip(" |"):
                parts.append(row_str)
    return "\n".join(parts)

def extract_pdf(path):
    import pdfplumber
    parts = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"\n--- Page {i+1} ---\n{text}")
            for table in (page.extract_tables() or []):
                parts.append("\n[table]")
                for row in table:
                    parts.append(" | ".join(str(c) if c else "" for c in row))
    return "\n".join(parts)

def extract_file(path):
    suffix = Path(path).suffix.lower()
    try:
        if suffix == ".docx":
            return extract_docx(path)
        elif suffix == ".doc":
            return extract_doc(path)
        elif suffix == ".xlsx":
            return extract_xlsx(path)
        elif suffix == ".pdf":
            return extract_pdf(path)
        else:
            return f"[skipped: unsupported {suffix}]"
    except Exception as e:
        return f"[ERROR: {e}]"

def main():
    if len(sys.argv) < 3:
        print("usage: extract_tender.py <tender_folder> <output_file>")
        sys.exit(1)
    folder = Path(sys.argv[1])
    output = Path(sys.argv[2])

    files = []
    for ext in ("*.doc", "*.docx", "*.xlsx", "*.pdf"):
        files.extend(folder.rglob(ext))

    files = sorted(files)
    with open(output, "w", encoding="utf-8") as f:
        for path in files:
            rel = path.relative_to(folder)
            f.write(f"\n\n{'#'*80}\n# FILE: {rel}\n{'#'*80}\n\n")
            print(f"extracting {rel}", file=sys.stderr)
            try:
                f.write(extract_file(path))
            except Exception as e:
                f.write(f"[FATAL: {e}]")
    print(f"done -> {output}", file=sys.stderr)

if __name__ == "__main__":
    main()
