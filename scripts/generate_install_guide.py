"""
Генерирует «Руководство по установке Tender AI Tools» в формате .docx.

Использование:
    python scripts/generate_install_guide.py
    python scripts/generate_install_guide.py --output "Установка Tender AI Tools.docx"
"""

import argparse
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, Mm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import sys
    print("ERROR: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ── цвета ──────────────────────────────────────────────────────────────────
CLR_ACCENT   = RGBColor(0x1F, 0x63, 0xB5)   # синий для заголовков шагов
CLR_BG_STEP  = "EBF2FB"                       # фон карточки шага
CLR_BG_TIP   = "FFF8E1"                       # фон совета
CLR_BG_WARN  = "FDE8E8"                       # фон предупреждения
CLR_BG_CODE  = "F4F4F4"                       # фон команды терминала
CLR_TEXT_MUTED = RGBColor(0x55, 0x55, 0x55)  # серый текст


def _set_margins(doc):
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(1.5)
        sec.page_width    = Mm(210)
        sec.page_height   = Mm(297)


def _run(para, text, *, bold=False, italic=False, size=11,
         color=None, font="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:ascii"), font)
    rf.set(qn("w:hAnsi"), font)
    return run


def _para(doc, text="", *, align=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=0, space_after=6, bold=False, italic=False,
          size=11, color=None, font="Calibri"):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(14)
    if text:
        _run(p, text, bold=bold, italic=italic, size=size,
             color=color, font=font)
    return p


def _blank(doc, space=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space)


def _shaded_table(doc, text, bg_hex, *, text_color=None, size=10.5,
                  bold=False, italic=False, prefix="", prefix_color=None,
                  indent_cm=0):
    """Однострочная таблица с заливкой фона — имитация цветного блока."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # ширина колонки
    cell = tbl.rows[0].cells[0]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # заливка
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg_hex)
    tcPr.append(shd)

    # отступы внутри ячейки
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top","60"),("bottom","60"),("left","120"),("right","120")]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), val)
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

    # нет рамок
    bdr = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        bdr.append(el)
    tcPr.append(bdr)

    # текст
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if prefix:
        _run(p, prefix + "  ", bold=True, size=size+1,
             color=prefix_color or RGBColor(0,0,0))
    _run(p, text, bold=bold, italic=italic, size=size,
         color=text_color or RGBColor(0x22, 0x22, 0x22))

    _blank(doc, 6)
    return tbl


def _step_header(doc, number, title):
    """Большой заголовок шага: «Шаг 1 — Установить Python»."""
    _blank(doc, 10)
    p = _para(doc, space_before=0, space_after=4)
    _run(p, f"Шаг {number}", bold=True, size=18, color=CLR_ACCENT)
    _run(p, f"  —  {title}", bold=False, size=16,
         color=RGBColor(0x33, 0x33, 0x33))


def _section_title(doc, text):
    _blank(doc, 8)
    p = _para(doc, text, bold=True, size=14,
              color=RGBColor(0x1A, 0x1A, 0x1A), space_after=4)
    return p


def _code(doc, text):
    """Моноширинная строка команды в сером блоке."""
    _code_line(doc, text)


def _code_line(doc, text):
    """Моноширинная строка команды в сером блоке — через функцию."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), CLR_BG_CODE)
    tcPr.append(shd)

    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top","60"),("bottom","60"),("left","180"),("right","120")]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), val)
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

    bdr = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        bdr.append(el)
    tcPr.append(bdr)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:ascii"), "Courier New")
    rf.set(qn("w:hAnsi"), "Courier New")

    _blank(doc, 6)


def _tip(doc, text):
    _shaded_table(doc, text, CLR_BG_TIP, prefix=">>", size=10.5,
                  prefix_color=RGBColor(0xF5, 0x7F, 0x17))


def _warn(doc, text):
    _shaded_table(doc, text, CLR_BG_WARN, prefix="(!)", size=10.5,
                  text_color=RGBColor(0xB7, 0x1C, 0x1C),
                  prefix_color=RGBColor(0xB7, 0x1C, 0x1C))


def _bullet(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(3)
        _run(p, item, size=11)


def _numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(3)
        _run(p, item, size=11)


# ── страницы ────────────────────────────────────────────────────────────────

def page_title(doc):
    _blank(doc, 60)
    p = _para(doc, "Tender AI Tools",
              align=WD_ALIGN_PARAGRAPH.CENTER,
              bold=True, size=32, color=CLR_ACCENT, space_after=10)

    p = _para(doc, "Руководство по установке",
              align=WD_ALIGN_PARAGRAPH.CENTER,
              size=18, color=RGBColor(0x44, 0x44, 0x44), space_after=6)

    today = datetime.now().strftime("%d.%m.%Y")
    _para(doc, today,
          align=WD_ALIGN_PARAGRAPH.CENTER,
          size=12, color=CLR_TEXT_MUTED, space_after=0)

    doc.add_page_break()


def page_whats_inside(doc):
    _section_title(doc, "Что в архиве")

    _para(doc, "Архив содержит всё необходимое для работы:", size=11, space_after=4)

    from docx.shared import Pt, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    rows = [
        ("📁 skills/", "4 скилла для Cowork — основные инструменты анализа закупок"),
        ("📁 mcp/",     "MCP-сервер — скачивает документы с ЕИС прямо в чат"),
        ("📄 README.md","Краткий обзор проекта"),
    ]

    tbl = doc.add_table(rows=len(rows)+1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # заголовок
    hdr = tbl.rows[0].cells
    for i, txt in enumerate(["Папка / файл", "Назначение"]):
        p = hdr[i].paragraphs[0]
        _run(p, txt, bold=True, size=10.5)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9E4F5")
        tcPr.append(shd)

    for r, (folder, desc) in enumerate(rows, start=1):
        cells = tbl.rows[r].cells
        p0 = cells[0].paragraphs[0]
        _run(p0, folder, bold=True, size=10, font="Courier New")
        p1 = cells[1].paragraphs[0]
        _run(p1, desc, size=10.5)

    _blank(doc, 6)


def page_requirements(doc):
    _section_title(doc, "Что нужно установить заранее")

    items = [
        "Python 3.10 или новее — python.org → Downloads → Python 3.x.x",
        "Claude Code — desktopприложение Anthropic (получить ссылку у меня)",
        "Аккаунт Cowork — получить приглашение у меня",
        "API-ключ DaMIA — зарегистрироваться на damia.ru (тариф «API-Старт», бесплатно)",
    ]
    _bullet(doc, items)

    _tip(doc, "Python: при установке обязательно поставь галочку «Add Python to PATH» — иначе команды pip/python не будут работать в терминале.")
    _warn(doc, "API-ключ DaMIA — это личный ключ. Никому не передавай и не пиши в общедоступных местах.")


def page_step1(doc):
    _step_header(doc, 1, "Установить Python")

    _para(doc, "Перейди на python.org → раздел Downloads → скачай последнюю версию Python 3.x.", size=11, space_after=4)

    _numbered(doc, [
        "Запусти скачанный установщик (.exe).",
        'На первом экране ОБЯЗАТЕЛЬНО поставь галочку "Add Python 3.x to PATH".',
        'Нажми "Install Now".',
        "Дождись завершения установки.",
    ])

    _blank(doc, 4)
    _para(doc, "Проверка — открой Терминал (Win+R → cmd → Enter) и введи:", size=11, space_after=3)
    _code_line(doc, "python --version")
    _para(doc, "Должна появиться строка вида: Python 3.10.x", size=10.5,
          italic=True, color=CLR_TEXT_MUTED, space_after=0)


def page_step2(doc):
    _step_header(doc, 2, "Распаковать архив")

    _para(doc, "Распакуй архив tender-ai-tools.zip в удобное место.", size=11, space_after=4)

    _tip(doc, "Рекомендуем: C:\\ClaudeCode\\  — путь без пробелов и кириллицы, команды в терминале работают надёжнее.")

    _para(doc, "После распаковки должна появиться папка:", size=11, space_after=3)
    _code_line(doc, "C:\\ClaudeCode\\tender-ai-tools\\")
    _para(doc, "Внутри — папки skills/, mcp/ и файл README.md.", size=10.5,
          italic=True, color=CLR_TEXT_MUTED, space_after=0)


def page_step3(doc):
    _step_header(doc, 3, "Установить зависимости")

    _para(doc, "Открой Терминал: нажми Win+R, введи cmd, нажми Enter.", size=11, space_after=6)

    _para(doc, "Выполни по очереди три команды:", size=11, space_after=3)
    _code_line(doc, "cd C:\\ClaudeCode\\tender-ai-tools\\mcp")
    _code_line(doc, "pip install -r requirements.txt")
    _code_line(doc, "pip install python-docx")

    _tip(doc, "Если pip не найден — закрой терминал, переустанови Python с галочкой «Add to PATH», открой терминал заново.")
    _warn(doc, "Если появляется ошибка «Access denied» — запускай cmd от имени администратора (правая кнопка мыши на cmd → «Запустить от имени администратора»).")


def page_step4(doc):
    _step_header(doc, 4, "Настроить MCP-сервер")

    _para(doc, "MCP-сервер скачивает документы закупок с сайта ЕИС по реестровому номеру.", size=11, space_after=6)

    _para(doc, "4.1. Создай файл с API-ключом:", size=11, bold=True, space_after=3)
    _para(doc, "Зайди в папку  C:\\ClaudeCode\\tender-ai-tools\\mcp\\", size=11, space_after=3)
    _para(doc, "Найди файл .env.example. Скопируй его и переименуй копию в  .env  (без расширения .example).", size=11, space_after=3)

    _tip(doc, "В Windows по умолчанию скрыты расширения файлов. Если файл называется «.env.example» — просто скопируй и переименуй, убрав «.example» с конца.")

    _blank(doc, 4)
    _para(doc, "4.2. Открой файл .env в Блокноте:", size=11, bold=True, space_after=3)
    _para(doc, "Правая кнопка мыши на .env → Открыть с помощью → Блокнот.", size=11, space_after=3)
    _para(doc, "Замени текст  your_api_key_here  на свой ключ DaMIA:", size=11, space_after=3)
    _code_line(doc, "DAMIA_API_KEY=вставь_свой_ключ_здесь")
    _para(doc, "Сохрани файл (Ctrl+S) и закрой.", size=11, space_after=0)


def page_step5(doc):
    _step_header(doc, 5, "Подключить MCP к Claude Code")

    _para(doc, "Это даст возможность скачивать документы закупок прямо из чата.", size=11, space_after=6)

    _numbered(doc, [
        "Открой приложение Claude Code.",
        "Нажми на иконку настроек (⚙) в левом нижнем углу.",
        'Выбери раздел "MCP Servers".',
        'Нажми "Add server" (или кнопку + Add).',
    ])

    _blank(doc, 4)
    _para(doc, "Заполни поля:", size=11, space_after=3)

    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    rows = [
        ("Name",    "zakupki-eis"),
        ("Type",    "stdio"),
        ("Command", "python"),
        ("Args",    "C:\\ClaudeCode\\tender-ai-tools\\mcp\\server.py"),
    ]
    tbl = doc.add_table(rows=len(rows)+1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, txt in enumerate(["Поле", "Значение"]):
        p = hdr[i].paragraphs[0]
        _run(p, txt, bold=True, size=10.5)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9E4F5")
        tcPr.append(shd)
    for r, (field, val) in enumerate(rows, start=1):
        cells = tbl.rows[r].cells
        _run(cells[0].paragraphs[0], field, bold=True, size=10.5)
        _run(cells[1].paragraphs[0], val, size=10.5, font="Courier New")
    _blank(doc, 6)

    _numbered(doc, [
        'Нажми "Save" / "Add".',
        "Перезапусти Claude Code.",
    ])

    _blank(doc, 4)
    _para(doc, "Проверка:", size=11, bold=True, space_after=3)
    _para(doc, 'Создай новый чат в Claude Code. Напечатай @zakupki-eis — должен появиться список инструментов (download_tender, check_tender_updates и др.).', size=11, space_after=0)
    _tip(doc, "Если @zakupki-eis не появляется — перезапусти Claude Code ещё раз. Иногда нужно подождать 10–15 секунд после запуска.")


def page_step6(doc):
    _step_header(doc, 6, "Установить скиллы в Cowork")

    _para(doc, "Скиллы — это готовые инструкции для ИИ: анализ закупок, жалобы, запросы разъяснений.", size=11, space_after=6)

    _para(doc, "В папке  C:\\ClaudeCode\\tender-ai-tools\\skills\\  лежат 4 zip-файла:", size=11, space_after=3)

    skills = [
        ("verdikt-zakupki-v1.0.zip",       "Отчёт «берёмся / не берёмся» по документации закупки"),
        ("scan-zakupki-v0.3.zip",           "Детальный анализ ТЗ, смет, выявление расхождений"),
        ("zhaloba-fas-v0.1.zip",            "Жалоба в УФАС: нарушения, защита от РНП, сговор"),
        ("zapros-razyasneniy-v0.1.1.zip",   "Запрос разъяснений + генерация Word-файла"),
    ]
    for fname, desc in skills:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        _run(p, fname, bold=True, size=10, font="Courier New")
        _run(p, f"  — {desc}", size=10.5)

    _blank(doc, 6)
    _para(doc, "Для каждого zip-файла:", size=11, bold=True, space_after=3)
    _numbered(doc, [
        "Открой Cowork в браузере.",
        'Зайди в раздел "Skills" (или "Скиллы").',
        'Нажми кнопку "Upload" (или "Загрузить").',
        "Выбери zip-файл.",
        'Убедись, что скилл появился в списке со статусом "Active".',
    ])

    _warn(doc, "Важно: загружай именно zip-файл целиком, не распакованную папку. Cowork ожидает архив.")


def page_step7(doc):
    _step_header(doc, 7, "Проверить работу")

    _para(doc, "Всё готово. Проверяем:", size=11, space_after=6)

    _numbered(doc, [
        "Открой Claude Code.",
        "Создай новый чат.",
        "Прикрепи любой файл из закупки (ТЗ, смета, PDF из ЕИС).",
        "Напиши: оцени закупку",
    ])

    _blank(doc, 4)
    _para(doc, 'ИИ должен запустить скилл «verdikt-zakupki» и выдать отчёт «берёмся / не берёмся» с карточкой закупки.', size=11, space_after=6)

    _tip(doc, "Если скилл не запустился — убедись, что он в статусе Active в Cowork. Иногда нужно обновить страницу.")

    _blank(doc, 8)
    _para(doc, "Другие команды для проверки:", size=11, bold=True, space_after=3)

    cmds = [
        ("составь жалобу",          "→ запускает скилл zhaloba-fas"),
        ("сделай запрос разъяснений", "→ запускает скилл zapros-razyasneniy"),
    ]
    for cmd, hint in cmds:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        _run(p, f'«{cmd}»', bold=True, size=11)
        _run(p, f"  {hint}", size=11, color=CLR_TEXT_MUTED)


def page_contacts(doc):
    _blank(doc, 20)
    doc.add_page_break()

    _blank(doc, 60)
    _para(doc, "Вопросы?", align=WD_ALIGN_PARAGRAPH.CENTER,
          bold=True, size=20, color=CLR_ACCENT, space_after=10)
    _para(doc, "Пиши мне — помогу разобраться.",
          align=WD_ALIGN_PARAGRAPH.CENTER, size=14,
          color=RGBColor(0x44, 0x44, 0x44), space_after=0)


# ── сборка ───────────────────────────────────────────────────────────────────

def generate(output_path: Path):
    doc = Document()
    _set_margins(doc)

    page_title(doc)
    page_whats_inside(doc)
    page_requirements(doc)
    page_step1(doc)
    page_step2(doc)
    page_step3(doc)
    page_step4(doc)
    page_step5(doc)
    page_step6(doc)
    page_step7(doc)
    page_contacts(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"OK: {output_path}  ({output_path.stat().st_size // 1024} KB)")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", "-o", type=Path,
                        default=Path("Установка Tender AI Tools.docx"))
    args = parser.parse_args()
    generate(args.output)


if __name__ == "__main__":
    main()
